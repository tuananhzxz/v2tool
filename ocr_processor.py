import os
from PIL import Image, ImageEnhance
import io
from google import genai
from docx import Document
import time

class OCRProcessor:
    def __init__(self, api_key):
        self.api_key = api_key
        self.client = genai.Client(api_key=api_key)
        
    def process_image(self, image_path):
        """Xử lý một ảnh và trả về văn bản đã nhận dạng"""
        try:
            # Đọc và xử lý ảnh
            with Image.open(image_path) as img:
                # Kiểm tra kích thước ảnh
                max_size = (1920, 1080)
                if img.size[0] > max_size[0] or img.size[1] > max_size[1]:
                    img.thumbnail(max_size, Image.Resampling.LANCZOS)
                
                # Chuyển sang grayscale và tăng độ tương phản
                img = img.convert('L')
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(1.5)
                
                # Chuyển ảnh đã xử lý thành bytes
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_byte_arr = img_byte_arr.getvalue()
                
                # Tạo prompt cho việc nhận dạng văn bản
                prompt = f"""
                    NHIỆM VỤ: Nhận dạng và trích xuất văn bản từ ảnh với độ chính xác cao nhất.
                    
                    YÊU CẦU CHẤT LƯỢNG:
                    1. Nhận dạng chính xác 100% nội dung văn bản, kể cả chữ nhỏ
                    2. Phân biệt rõ các đoạn văn bản khác nhau, các bóng thoại khác nhau
                    3. Giữ nguyên vị trí và thứ tự của các bóng thoại
                    4. Không bỏ sót bất kỳ ký tự nào
                    
                    QUY TẮC XỬ LÝ:
                    1. Loại bỏ các yếu tố không phải văn bản 
                    2. QUAN TRỌNG: Xử lý mỗi bóng thoại (speech bubble) như MỘT CÂU HOÀN CHỈNH TRÊN MỘT DÒNG DUY NHẤT
                    3. Mỗi bóng thoại riêng biệt sẽ được xuất ra thành một dòng văn bản riêng biệt
                    4. Giữ nguyên các dấu câu và định dạng đặc biệt
                    
                    ĐỊNH DẠNG ĐẦU RA:
                    - Mỗi bóng thoại trên một dòng riêng
                    - Giữ nguyên các dấu câu và định dạng
                    - Không thêm bất kỳ chú thích hay giải thích nào
                    - Không tách văn bản trong một bóng thoại thành nhiều dòng
                """
                
                # Gửi yêu cầu đến Gemini API với cấu hình tối ưu
                response = self.client.models.generate_content(
                    model="gemini-2.5-pro-exp-03-25",
                    contents=[
                        {"text": prompt},
                        {"inline_data": {"mime_type": "image/png", "data": img_byte_arr}}
                    ],
                    config={
                        "temperature": 0.1,  # Giảm temperature để tăng độ chính xác
                        "max_output_tokens": 2048,  # Tăng max tokens để xử lý văn bản dài
                        "top_p": 0.8,
                        "top_k": 40,
                        "candidate_count": 1
                    }
                )
                
                if response.text:
                    text = response.text.strip()
                    print(f"Đã nhận dạng văn bản từ file {image_path}: {text[:100]}...")  # Log để debug
                    return text
                else:
                    print(f"Không nhận dạng được văn bản từ file {image_path}")
                    return None
                    
        except Exception as e:
            print(f"Lỗi khi xử lý file ảnh {image_path}: {str(e)}")
            return None
            
    def process_files(self, files, output_path):
        """Xử lý nhiều file và lưu kết quả vào file Word"""
        try:
            # Tạo tài liệu Word mới
            doc = Document()
            processed_files = []
            failed_files = []
            
            print(f"Bắt đầu xử lý {len(files)} files...")
            
            # Xử lý từng file
            for file in files:
                if not file.filename:
                    continue
                    
                try:
                    # Kiểm tra định dạng file
                    file_ext = os.path.splitext(file.filename)[1].lower()
                    
                    if file_ext in ['.jpg', '.jpeg', '.png', '.webp']:
                        print(f"Đang xử lý file ảnh: {file.filename}")
                        # Lưu file ảnh tạm thời
                        temp_path = os.path.join(os.path.dirname(output_path), file.filename)
                        file.save(temp_path)
                        
                        # Thêm thông tin về file đang xử lý
                        doc.add_paragraph(f"=== File: {file.filename} ===")
                        
                        # Xử lý ảnh
                        text = self.process_image(temp_path)
                        if text:
                            # Tách văn bản thành các dòng
                            lines = text.split('\n')
                            for line in lines:
                                if line.strip():  # Chỉ thêm dòng không trống
                                    doc.add_paragraph(line.strip())
                            processed_files.append(file.filename)
                            print(f"Đã xử lý thành công file: {file.filename}")
                        else:
                            doc.add_paragraph("Không nhận dạng được văn bản từ ảnh này.")
                            failed_files.append(file.filename)
                            print(f"Không thể nhận dạng văn bản từ file: {file.filename}")
                            
                        doc.add_paragraph('---')  # Thêm dấu phân cách giữa các file
                        
                        # Xóa file tạm
                        os.remove(temp_path)
                        
                        # Thêm delay giữa các lần xử lý
                        time.sleep(1)
                        
                    else:
                        failed_files.append(file.filename)
                        print(f"File không được hỗ trợ: {file.filename}")
                        continue
                        
                except Exception as e:
                    print(f"Lỗi khi xử lý file {file.filename}: {str(e)}")
                    failed_files.append(file.filename)
                    continue
            
            if not processed_files:
                print("Không có file nào được xử lý thành công")
                return False, processed_files, failed_files
                
            # Lưu tài liệu
            print(f"Đang lưu file Word tại: {output_path}")
            doc.save(output_path)
            
            # Kiểm tra xem file Word có nội dung không
            if os.path.getsize(output_path) == 0:
                print("File Word trống sau khi lưu")
                return False, processed_files, failed_files
                
            print(f"Đã lưu file Word thành công với {len(processed_files)} files đã xử lý")
            return True, processed_files, failed_files
            
        except Exception as e:
            print(f"Lỗi khi xử lý files: {str(e)}")
            return False, [], files

def process_folder(folder_path, api_key=None, genre=None, target_langs=None):
    """Xử lý tất cả ảnh trong thư mục và tạo file Word cho mỗi ngôn ngữ"""
    try:
        processor = OCRProcessor(api_key)
        
        # Lấy danh sách ảnh
        image_files = [f for f in os.listdir(folder_path) if
                      f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp'))]
        
        if not image_files:
            raise Exception("Không tìm thấy ảnh trong thư mục")
            
        # Sắp xếp ảnh theo số trong tên file
        image_files.sort(key=lambda x: int(re.search(r'\d+', x).group() if re.search(r'\d+', x) else 0))
        
        # Ngôn ngữ mặc định nếu không được chỉ định
        if not target_langs:
            target_langs = ['vi']
        
        # Xử lý từng ảnh
        output_files = {lang: [] for lang in target_langs}
        for idx, img_file in enumerate(image_files, 1):
            img_path = os.path.join(folder_path, img_file)
            
            # Thực hiện OCR
            text = processor.process_image(img_path)
            
            if text:
                # Dịch và lưu cho từng ngôn ngữ
                for lang in target_langs:
                    # Dịch văn bản nếu có API key
                    translated_text = text
                    if api_key:
                        translated_text = processor.translate_text(text, lang, genre)
                    
                    # Tạo tên file Word theo định dạng số thứ tự
                    word_filename = f"{idx}_{lang}.docx"
                    output_path = os.path.join(folder_path, word_filename)
                    
                    # Lưu vào file Word
                    if processor.save_to_word(translated_text, output_path):
                        output_files[lang].append(output_path)
        
        if not any(output_files.values()):
            raise Exception("Không thể trích xuất được văn bản từ các ảnh")
        
        return output_files
        
    except Exception as e:
        print(f"Lỗi trong quá trình xử lý OCR: {str(e)}")
        raise 