import os
import glob
from docx import Document
import re


def natural_sort_key(s):
    """Hàm sắp xếp tập tin theo số tự nhiên (01, 02, 03,... thay vì 1, 10, 2)"""
    return [int(text) if text.isdigit() else text.lower()
            for text in re.split(r'(\d+)', s)]


def merge_word_documents(input_folder, output_file):
    """
    Gộp tất cả tập tin Word từ thư mục vào một tập tin duy nhất

    Args:
        input_folder: Đường dẫn đến thư mục chứa các tập tin Word
        output_file: Đường dẫn đến tập tin kết quả
    """
    # Tạo document mới
    merged_document = Document()

    # Tìm tất cả tập tin .docx trong thư mục
    file_pattern = os.path.join(input_folder, "*.docx")
    files = glob.glob(file_pattern)

    # Sắp xếp tập tin theo thứ tự số tự nhiên
    files.sort(key=natural_sort_key)

    if not files:
        print(f"Không tìm thấy tập tin Word nào trong thư mục '{input_folder}'")
        return False

    # Đếm số lượng tập tin đã xử lý
    processed_files = 0

    # Gộp các tập tin
    for file_path in files:
        file_name = os.path.basename(file_path)

        # Kiểm tra tập tin có chứa "_converted" không
        if "_converted" in file_name:
            try:
                print(f"Đang xử lý: {file_name}")

                # Mở tập tin
                doc = Document(file_path)

                # Sao chép tất cả các đoạn
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # Chỉ sao chép đoạn không trống
                        merged_document.add_paragraph(paragraph.text)

                processed_files += 16

            except Exception as e:
                print(f"Lỗi khi xử lý tập tin '{file_name}': {str(e)}")

    # Lưu tập tin kết quả
    try:
        merged_document.save(output_file)
        print(f"\nĐã gộp {processed_files} tập tin thành công!")
        print(f"Kết quả đã được lưu vào: {output_file}")
        return True
    except Exception as e:
        print(f"Lỗi khi lưu tập tin kết quả: {str(e)}")
        return False


# Hàm chính
def main():
    print("CÔNG CỤ GỘP TÀI LIỆU WORD")
    print("=========================")

    # Nhập đường dẫn thư mục đầu vào
    input_folder = input("Nhập đường dẫn đến thư mục chứa các tập tin Word: ")

    # Kiểm tra thư mục tồn tại
    if not os.path.isdir(input_folder):
        print(f"Lỗi: Thư mục '{input_folder}' không tồn tại!")
        return

    # Nhập tên tập tin đầu ra
    output_file = input("Nhập đường dẫn tập tin kết quả (mặc định: merged_document.docx): ")
    if not output_file:
        output_file = "merged_document.docx"
    elif not output_file.endswith(".docx"):
        output_file += ".docx"

    # Thực hiện gộp tài liệu
    merge_word_documents(input_folder, output_file)


if __name__ == "__main__":
    main()