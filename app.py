import time
from flask import Flask, render_template, request, jsonify, send_file, url_for, Response
import os
from werkzeug.utils import secure_filename
from PIL import Image, ImageEnhance
import shutil
from changeImage import convert_images
from cutmergeimage import ImageProcessor
from mergeWord import merge_word_documents
from renameImage import rename_files
import requests
from bs4 import BeautifulSoup
import io
import zipfile
import re
from flask_cors import CORS
from ocr_processor import process_folder
from add_logo import LogoProcessor
from docx import Document
from google import genai
from google.genai import types
from dotenv import load_dotenv
from service.dowloadImg import (
    download_selected_images,
)
import base64  # Add this import at the top with other imports
from ocr_processor import OCRProcessor

# Load environment variables
load_dotenv()

# Modify the app configuration to use in-memory processing
app = Flask(__name__, static_folder='service')
CORS(app)
app.config['UPLOAD_FOLDER'] = '/tmp'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max-limit

# Create temporary directory if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


# Hàm xử lý lỗi chung
def handle_error(e):
    error_msg = str(e)
    if isinstance(e, requests.exceptions.RequestException):
        error_msg = "Lỗi kết nối: Không thể tải dữ liệu từ URL"
    elif isinstance(e, IOError):
        error_msg = "Lỗi đọc/ghi file"
    elif isinstance(e, Image.UnidentifiedImageError):
        error_msg = "File ảnh không hợp lệ hoặc bị hỏng"
    return jsonify({'error': error_msg})

@app.route('/')
def index():
    functions = {
        'changeImage': 'Chuyển đổi định dạng ảnh',
        'cutmergeimage': 'Cắt và ghép ảnh',
        'dowloaf': 'Tải ảnh từ web',
        'mergeWord': 'Gộp file Word',
        'renameImage': 'Đổi tên ảnh',
        'ocr': 'Nhận dạng chữ trong ảnh (OCR)',
        'addlogo': 'Thêm logo vào ảnh'
    }
    return render_template('index.html', functions=functions)

@app.route('/function/<function_name>')
def function_page(function_name):
    templates = {
        'changeImage': 'functions/changeimage.html',
        'cutmergeimage': 'functions/cutmergeimage.html',
        'dowloaf': 'functions/dowloaf.html',
        'mergeWord': 'functions/mergeword.html',
        'renameImage': 'functions/renameimage.html',
        'ocr': 'functions/ocr.html',
        'translate': 'functions/translate.html',
        'addlogo': 'functions/addlogo.html'  # Add this line
    }
    
    if function_name not in templates:
        return jsonify({'error': 'Chức năng không tồn tại'}), 404
        
    return render_template(templates[function_name])

@app.route('/execute/changeImage', methods=['POST'])
def execute_change_image():
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'Không có file được tải lên'})
        
        files = request.files.getlist('files')
        target_format = request.form.get('target_format', 'webp')
        
        if not files or files[0].filename == '':
            return jsonify({'error': 'Không có file được chọn'})
            
        # Tạo thư mục tạm thời
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_convert')
        output_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'output_convert')
        
        for dir_path in [temp_dir, output_dir]:
            if os.path.exists(dir_path):
                shutil.rmtree(dir_path)
            os.makedirs(dir_path)
        
        # Theo dõi các chuyển đổi đã thực hiện
        conversions = []
            
        # Lưu và xử lý từng file
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
                
                # Xác định định dạng nguồn và chuyển đổi
                with Image.open(file_path) as img:
                    source_format = img.format.lower() if img.format else os.path.splitext(filename)[1][1:].lower()
                    # Chuyển đổi và lưu với định dạng mới
                    output_path = os.path.join(output_dir, os.path.splitext(filename)[0] + '.' + target_format.lower())
                    
                    # Thêm thông tin chuyển đổi
                    conversions.append(f"{source_format.upper()} → {target_format.upper()}")
                    
                    if target_format.upper() == 'WEBP':
                        img.save(output_path, 'WEBP', quality=90)
                    elif target_format.upper() == 'JPEG':
                        if img.mode in ('RGBA', 'LA'):
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            background.paste(img, mask=img.split()[-1])
                            background.save(output_path, 'JPEG', quality=95)
                        else:
                            img.save(output_path, 'JPEG', quality=95)
                    else:  # PNG
                        img.save(output_path, 'PNG')
        
        # Tạo file zip
        zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'converted_images.zip')
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, dirs, files in os.walk(output_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, output_dir)
                    zipf.write(file_path, arcname)
        
        # Read the zip file into memory
        with open(zip_path, 'rb') as f:
            zip_data = f.read()
        
        # Encode the zip data as base64 for JSON transmission
        zip_base64 = base64.b64encode(zip_data).decode('utf-8')
        
        # Xóa thư mục tạm
        shutil.rmtree(temp_dir)
        shutil.rmtree(output_dir)
        os.remove(zip_path)
        
        # Tạo thông báo chi tiết về các chuyển đổi
        unique_conversions = list(set(conversions))
        conversion_message = ", ".join(unique_conversions)
        
        return jsonify({
            'message': f'Chuyển đổi thành công! ({conversion_message})',
            'zip_data': zip_base64,
            'filename': 'converted_images.zip',
            'conversions': unique_conversions
        })
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/cutmergeimage', methods=['POST'])
def execute_cut_merge_image():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    action = request.form.get('action')
    min_height = int(request.form.get('min_height', 2500))  # Chiều cao tối thiểu mặc định 2500px
    images_per_group = int(request.form.get('parts', 2))  # Số ảnh mỗi nhóm khi ghép
    width = request.form.get('width')
    height = request.form.get('height')
    
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
        
    # Tạo thư mục tạm thời
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_process')
    output_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'output_process')
    
    for dir_path in [temp_dir, output_dir]:
        if os.path.exists(dir_path):
            shutil.rmtree(dir_path)
        os.makedirs(dir_path)
        
    try:
        # Lưu các file
        saved_files = []
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
                saved_files.append(file_path)
                
                # Resize ảnh nếu có yêu cầu
                if width or height:
                    with Image.open(file_path) as img:
                        new_width = int(width) if width else img.width
                        new_height = int(height) if height else img.height
                        resized = img.resize((new_width, new_height))
                        resized.save(file_path)
        
        processor = ImageProcessor(temp_dir)
        
        try:
            if action == 'split':
                result_files = processor.split_images(min_height)
            else:  # merge
                if len(saved_files) < images_per_group:
                    return jsonify({'error': f'Số lượng ảnh ({len(saved_files)}) phải lớn hơn hoặc bằng số ảnh mỗi nhóm ({images_per_group})'})
                result_files = processor.combine_images(images_per_group)
            
            if not result_files:
                return jsonify({'error': 'Không thể xử lý ảnh. Vui lòng kiểm tra lại các tham số.'})
            
            # Copy các file kết quả vào thư mục output với tên mới
            for i, file_path in enumerate(result_files, 1):
                file_ext = os.path.splitext(file_path)[1]
                new_name = f"{i}{file_ext}"
                new_path = os.path.join(output_dir, new_name)
                shutil.copy2(file_path, new_path)
            
            # Tạo file zip chứa kết quả
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'processed_images.zip')
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for root, dirs, files in os.walk(output_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, output_dir)
                        zipf.write(file_path, arcname)
            
            # Xóa thư mục tạm
            shutil.rmtree(temp_dir)
            shutil.rmtree(output_dir)
            
            return jsonify({
                'message': 'Xử lý thành công!',
                'output_files': ['processed_images.zip']
            })
            
        except Exception as e:
            return jsonify({'error': str(e)})
            
    except Exception as e:
        return handle_error(e)

@app.route('/execute/mergeWord', methods=['POST'])
def execute_merge_word():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    # Tạo thư mục tạm thời
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_word')
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    try:
        # Lưu các file
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
        
        # Gộp file
        output_path = os.path.join(temp_dir, 'merged_document.docx')
        success = merge_word_documents(temp_dir, output_path)
        
        if success:
            # Tạo file zip chứa kết quả
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'merged_documents.zip')
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                zipf.write(output_path, 'merged_document.docx')
            
            # Xóa thư mục tạm
            shutil.rmtree(temp_dir)
            
            return jsonify({
                'message': 'Gộp file thành công!',
                'output_files': ['merged_documents.zip']
            })
        else:
            return jsonify({'error': 'Không thể gộp các file'})
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/renameImage', methods=['POST'])
def execute_rename_image():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    # Tạo thư mục tạm thời
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_rename')
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    try:
        # Lưu các file
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
        
        # Đổi tên file
        rename_files(temp_dir)
        
        # Tạo file zip chứa kết quả
        zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'renamed_images.zip')
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)
        
        # Xóa thư mục tạm
        shutil.rmtree(temp_dir)
        
        return jsonify({
            'message': 'Đổi tên file thành công!',
            'output_files': ['renamed_images.zip']
        })
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/ocr', methods=['POST'])
def execute_ocr():
    try:
        files = request.files.getlist('files')
        api_key = request.form.get('api_key')
        mode = request.form.get('mode', 'ocr')
        
        if not files or not api_key:
            return jsonify({'error': 'Vui lòng cung cấp file và API key'})
        
        # Tạo thư mục tạm thời
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_ocr')
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir)
        
        try:
            # Khởi tạo OCR processor
            processor = OCRProcessor(api_key)
            
            # Xử lý files và lưu kết quả
            output_filename = f'VBCĐ_{int(time.time())}.docx'
            output_path = os.path.join(temp_dir, output_filename)
            
            print(f"Bắt đầu xử lý OCR với {len(files)} files...")
            success, processed_files, failed_files = processor.process_files(files, output_path)
            
            if not success:
                print("Không thể xử lý bất kỳ file nào")
                return jsonify({'error': 'Không thể xử lý bất kỳ file nào'})
            
            print(f"Đã xử lý thành công {len(processed_files)} files")
            
            # Tạo file zip chứa kết quả
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'processed_ocr.zip')
            print(f"Đang tạo file ZIP tại: {zip_path}")
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                zipf.write(output_path, output_filename)
            
            # Read the zip file into memory
            with open(zip_path, 'rb') as f:
                zip_data = f.read()
            
            # Encode the zip data as base64 for JSON transmission
            zip_base64 = base64.b64encode(zip_data).decode('utf-8')
            
            print("Đã mã hóa file ZIP thành base64")
            
            # Xóa thư mục tạm và file zip
            shutil.rmtree(temp_dir)
            os.remove(zip_path)
            
            return jsonify({
                'success': True,
                'message': 'Xử lý thành công!',
                'zip_data': zip_base64,
                'filename': 'processed_ocr.zip',
                'processed_files': processed_files,
                'failed_files': failed_files
            })
            
        except Exception as e:
            print(f"Lỗi trong quá trình xử lý OCR: {str(e)}")
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            return jsonify({'error': str(e)})
        
    except Exception as e:
        print(f"Lỗi chung: {str(e)}")
        return jsonify({'error': str(e)})

def process_word_file(file_path, doc, mode, genres, styles, api_key, target_langs):
    try:
        print(f"Bắt đầu xử lý file Word: {file_path}")
        
        # Đọc file Word
        docx_doc = Document(file_path)
        text = ''
        
        # Đọc tất cả các đoạn văn bản
        for para in docx_doc.paragraphs:
            if para.text.strip():  # Chỉ lấy các đoạn không trống
                text += para.text + '\n'
        
        print(f"Đã đọc được {len(text.split('\n'))} đoạn văn bản")
        
        if text.strip():
            # Thêm thông tin về file đang xử lý
            doc.add_paragraph(f"=== File: {os.path.basename(file_path)} ===")
            
            if mode == 'translate':
                print("Bắt đầu quá trình dịch...")
                print(f"Ngôn ngữ đích: {target_langs}")
                
                for lang in target_langs:
                    print(f"Đang dịch sang {getLangName(lang)}...")
                    prompt = f"""
                        NHIỆM VỤ: Dịch văn bản sau sang {getLangName(lang)} với độ chính xác cao nhất.
                        
                        THỂ LOẠI: {', '.join(genres) if genres else 'Phù hợp với nội dung'}
                        PHONG CÁCH: {', '.join(styles) if styles else 'Phù hợp với nội dung'}
                        
                        VĂN BẢN CẦN DỊCH:
                        {text}
                    """
                    
                    client = genai.Client(api_key=api_key)
                    response = client.models.generate_content(
                        model="gemini-2.0-flash",
                        contents=prompt
                    )
                                                            
                    if response.text:
                        
                        # Tách văn bản dịch thành các đoạn và thêm vào document
                        translated_paragraphs = response.text.split('\n')
                        for para_text in translated_paragraphs:
                            if para_text.strip():  # Chỉ thêm các đoạn không trống
                                doc.add_paragraph(para_text.strip())
                        
                        # Thêm dấu phân cách giữa các bản dịch
                        doc.add_paragraph('---')
                        
                        # Thêm delay giữa các lần dịch
                        time.sleep(1)
                    else:
                        print("API trả về văn bản trống")
                        doc.add_paragraph(f"Không thể dịch sang {getLangName(lang)}")
            else:
                # Nếu không phải chế độ dịch, chỉ thêm văn bản gốc
                doc.add_paragraph(text)
                doc.add_paragraph('---')
                
            print("Đã xử lý xong file Word")
            return True
        else:
            print(f"File {file_path} không có nội dung văn bản")
            return False
            
    except Exception as e:
        print(f"Lỗi khi xử lý file Word {file_path}: {str(e)}")
        return False

def process_image_file(file_path, doc, mode, genres, styles, api_key):
    try:
        # Đọc và xử lý ảnh
        with Image.open(file_path) as img:
            # Kiểm tra kích thước ảnh
            max_size = (1920, 1080)
            if img.size[0] > max_size[0] or img.size[1] > max_size[1]:
                img.thumbnail(max_size, Image.Resampling.LANCZOS)
            
            # Chuyển sang grayscale và tăng độ tương phản
            img = img.convert('L')
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.5)
            
            # Chuyển ảnh đã xử lý thành bytes
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_byte_arr = img_byte_arr.getvalue()
            
            # Tạo prompt cho việc nhận dạng văn bản
            prompt = f"""
                NHIỆM VỤ: Nhận dạng và trích xuất văn bản từ ảnh với độ chính xác cao nhất.
                
                YÊU CẦU CHẤT LƯỢNG:
                1. Nhận dạng chính xác 100% nội dung văn bản, kể cả chữ nhỏ
                2. Phân biệt rõ các đoạn văn bản khác nhau, các bóng thoại khác nhau
                3. Giữ nguyên vị trí và thứ tự của các bóng thoại
                4. Không bỏ sót bất kỳ ký tự nào
                
                QUY TẮC XỬ LÝ:
                1. Loại bỏ các yếu tố không phải văn bản 
                2. QUAN TRỌNG: Xử lý mỗi bóng thoại (speech bubble) như MỘT CÂU HOÀN CHỈNH TRÊN MỘT DÒNG DUY NHẤT
                3. Mỗi bóng thoại riêng biệt sẽ được xuất ra thành một dòng văn bản riêng biệt
                4. Giữ nguyên các dấu câu và định dạng đặc biệt
                
                ĐỊNH DẠNG ĐẦU RA:
                - Mỗi bóng thoại trên một dòng riêng
                - Giữ nguyên các dấu câu và định dạng
                - Không thêm bất kỳ chú thích hay giải thích nào
                - Không tách văn bản trong một bóng thoại thành nhiều dòng
            """
            
            client = genai.Client(api_key=api_key)
            
            # Gửi yêu cầu đến Gemini API với cấu hình tối ưu
            response = client.models.generate_content(
                model="gemini-2.5-pro-exp-03-25",
                contents=[
                    {"text": prompt},
                    {"inline_data": {"mime_type": "image/png", "data": img_byte_arr}}
                ],
                config={
                    "temperature": 0.1,  # Giảm temperature để tăng độ chính xác
                    "max_output_tokens": 2048,  # Tăng max tokens để xử lý văn bản dài
                    "top_p": 0.8,
                    "top_k": 40,
                    "candidate_count": 1
                }
            )
            
            if response.text:
                # Thêm thông tin về file đang xử lý
                doc.add_paragraph(f"=== File: {os.path.basename(file_path)} ===")
                
                if mode == 'translate':
                    # Dịch văn bản
                    target_langs = request.form.getlist('target_langs[]')
                    for lang in target_langs:
                        prompt = f"""
                            Hãy dịch đoạn văn bản sau sang {getLangName(lang)}.
                            Thể loại: {', '.join(genres) if genres else 'Không xác định'}
                            Phong cách: {', '.join(styles) if styles else 'Không xác định'}
                            
                            Văn bản cần dịch:
                            {response.text}
                        """
                        
                        response = client.models.generate_content(
                            model="gemini-2.5-pro-exp-03-25",
                            contents=prompt,
                            config={
                                "temperature": 0.2,
                                "max_output_tokens": 2048
                            }
                        )
                        
                        if response.text:
                            doc.add_paragraph(f"=== {getLangName(lang)} ===")
                            doc.add_paragraph(response.text)
                            doc.add_paragraph('---')
                else:
                    doc.add_paragraph(response.text)
                    doc.add_paragraph('---')
                    
                # Thêm delay giữa các lần xử lý để tránh rate limit
                time.sleep(1)
                
            return True
    except Exception as e:
        print(f"Lỗi khi xử lý file ảnh {file_path}: {str(e)}")
        return False

# Hàm chuyển đổi mã ngôn ngữ thành tên
def getLangName(code):
    langs = {
        'vie': 'Tiếng Việt',
        'eng': 'Tiếng Anh',
        'jpn': 'Tiếng Nhật',
        'kor': 'Tiếng Hàn'
    }
    return langs.get(code, code)

@app.route('/execute/merge_ocr', methods=['POST'])
def execute_merge_ocr():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    # Tạo thư mục tạm thời
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_merge_ocr')
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    try:
        # Lưu các file Word
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
        
        # Gộp các file Word
        output_path = os.path.join(temp_dir, 'merged_ocr.docx')
        success = merge_word_documents(temp_dir, output_path)
        
        if success:
            # Tạo file zip chứa kết quả
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'merged_ocr.zip')
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                zipf.write(output_path, 'merged_ocr.docx')
            
            # Xóa thư mục tạm
            shutil.rmtree(temp_dir)
            
            return jsonify({
                'message': 'Gộp file thành công!',
                'output_files': ['merged_ocr.zip']
            })
        else:
            return jsonify({'error': 'Không thể gộp các file'})
        
    except Exception as e:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        return handle_error(e)

# Modify the download route to support direct blob response
@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(file_path):
            return jsonify({'error': 'File không tồn tại'})
        
        # Read the file into memory
        with open(file_path, 'rb') as f:
            file_data = f.read()
        
        # Create a response with the file data
        response = Response(file_data)
        response.headers['Content-Type'] = 'application/octet-stream'
        response.headers['Content-Disposition'] = f'attachment; filename={filename}'
        
        # Delete the file after sending
        os.remove(file_path)
        
        return response
    except Exception as e:
        return handle_error(e)

# Add a new route for client-side processing
@app.route('/api/process_result', methods=['POST'])
def process_result():
    try:
        # This endpoint will be called by the client to notify that processing is complete
        # and the result is stored in IndexedDB
        data = request.json
        file_id = data.get('file_id')
        filename = data.get('filename')
        
        return jsonify({
            'success': True,
            'message': 'Xử lý thành công!',
            'file_id': file_id,
            'filename': filename
        })
    except Exception as e:
        return jsonify({'error': str(e)})

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File quá lớn. Kích thước tối đa là 16MB'}), 413

@app.errorhandler(500)
def internal_error(e):
    return jsonify({'error': 'Lỗi máy chủ nội bộ'}), 500

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Không tìm thấy trang'}), 404


@app.route('/changeImage')
def change_image():
    return render_template('functions/changeimage.html')

@app.route('/cutmergeimage')
def cut_merge_image():
    return render_template('functions/cutmergeimage.html')

@app.route('/dowloaf')
def download_image():
    return render_template('functions/dowloaf.html')

@app.route('/mergeword')
def merge_word():
    return render_template('functions/mergeword.html')

@app.route('/renameimage')
def rename_image():
    return render_template('functions/renameimage.html')

@app.route('/ocr')
def ocr():
    return render_template('functions/ocr.html')

@app.route('/translate')
def translate():
    return render_template('functions/translate.html')

@app.route('/remove_speech_bubbles')
def remove_speech_bubbles():
    return render_template('functions/speechbubble.html')

@app.route('/test_api_key', methods=['POST'])
def test_api_key():
    try:
        data = request.json
        api_key = data.get('api_key')
        
        if not api_key:
            return jsonify({'error': 'API key không được cung cấp'})
        
        client = genai.Client(api_key=api_key)
        
        # Tạo một tin nhắn đơn giản để kiểm tra API key
        response = client.models.generate_content(
            model="gemini-2.5-pro-exp-03-25", contents="Explain how AI works in a few words"
        )
        
        return jsonify({'success': True, 'message': 'API key hợp lệ'})
        
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/execute/download_images', methods=['POST'])
def execute_download():
    return download_selected_images()

@app.route('/add_logo', methods=['GET', 'POST'])
def add_logo():
    if request.method == 'GET':
        return render_template('functions/addlogo.html')
    
    try:
        # Kiểm tra file upload
        if 'files' not in request.files or 'logo' not in request.files:
            return jsonify({'error': 'Vui lòng chọn cả logo và ảnh cần xử lý'}), 400
            
        files = request.files.getlist('files')
        logo = request.files['logo']
        
        if not files or not logo or files[0].filename == '' or logo.filename == '':
            return jsonify({'error': 'Vui lòng chọn file'}), 400
            
        # Lấy các tham số
        position = request.form.get('position', 'top_left')
        scale = float(request.form.get('scale', 10)) / 100  # Chuyển đổi từ phần trăm sang thập phân
        
        # Tạo thư mục tạm thời
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_logo')
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir)
            
        try:
            # Lưu logo
            logo_path = os.path.join(temp_dir, 'logo.png')
            logo.save(logo_path)
            
            # Lưu các file ảnh
            for file in files:
                if file.filename:
                    filename = secure_filename(file.filename)
                    file_path = os.path.join(temp_dir, filename)
                    file.save(file_path)
                    
            # Xử lý ảnh
            processor = LogoProcessor()
            output_path = processor.process_folder(temp_dir, logo_path, position, scale)
            
            if output_path:
                # Tạo file zip chứa kết quả
                zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'processed_images.zip')
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    zipf.write(output_path, os.path.basename(output_path))
                
                # Xóa thư mục tạm
                shutil.rmtree(temp_dir)
                        
                return jsonify({
                    'success': True,
                    'message': 'Thêm logo thành công',
                    'download_url': url_for('download_file', filename='processed_images.zip', _external=True)
                })
            else:
                return jsonify({'error': 'Không thể xử lý ảnh'}), 500
            
        except Exception as e:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            return jsonify({'error': str(e)}), 500
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/addlogo')
def add_logo_page():
    return render_template('functions/addlogo.html')

@app.route('/service/IndexedDB.js')
def serve_indexeddb_js():
    return send_file('service/IndexedDB.js', mimetype='application/javascript')

@app.route('/service/fileProcessor.js')
def serve_fileprocessor_js():
    return send_file('service/fileProcessor.js', mimetype='application/javascript')

@app.route('/execute/translate', methods=['POST'])
def execute_translate():
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'Không có file được tải lên'})
            
        files = request.files.getlist('files')
        api_key = request.form.get('api_key')
        
        # Fix: Get target languages properly from the form
        # The issue might be with how Flask processes the multiple select
        target_langs = request.form.getlist('target_langs[]')
        
        # Debug output to see what's being received
        print(f"Received target languages: {target_langs}")
        
        # If target_langs is empty, try alternative ways to get it
        if not target_langs:
            # Try without the brackets
            target_langs = request.form.getlist('target_langs')
            print(f"Alternative method: {target_langs}")
            
            # If still empty, try getting as a single value and splitting
            if not target_langs:
                target_lang_str = request.form.get('target_langs[]')
                if target_lang_str:
                    target_langs = [lang.strip() for lang in target_lang_str.split(',')]
                    print(f"Split method: {target_langs}")
        
        genres = request.form.getlist('genres[]')
        styles = request.form.getlist('styles[]')
        
        print(f"Bắt đầu xử lý dịch với {len(files)} files")
        print(f"Ngôn ngữ đích: {target_langs}")
        print(f"Thể loại: {genres}")
        print(f"Phong cách: {styles}")
        
        if not files or not api_key:
            return jsonify({'error': 'Vui lòng cung cấp file và API key'})
            
        if not target_langs:
            return jsonify({'error': 'Vui lòng chọn ít nhất một ngôn ngữ đích'})
            
        # Tạo thư mục tạm thời
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_translate')
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir)
        
        try:
            # Tạo tài liệu Word mới
            doc = Document()
            processed_files = []
            failed_files = []
            
            # Xử lý từng file
            for file in files:
                if not file.filename:
                    continue
                    
                try:
                    # Kiểm tra định dạng file
                    file_ext = os.path.splitext(file.filename)[1].lower()
                    print(f"Đang xử lý file: {file.filename} ({file_ext})")
                    
                    if file_ext in ['.docx', '.doc']:
                        # Lưu file Word tạm thời
                        temp_path = os.path.join(temp_dir, file.filename)
                        file.save(temp_path)
                        print(f"Đã lưu file tạm tại: {temp_path}")
                        
                        # Xử lý file Word
                        success = process_word_file(temp_path, doc, 'translate', genres, styles, api_key, target_langs)
                        if success:
                            processed_files.append(file.filename)
                            print(f"Đã xử lý thành công file: {file.filename}")
                        else:
                            failed_files.append(file.filename)
                            print(f"Không thể xử lý file: {file.filename}")
                            
                        # Xóa file tạm
                        os.remove(temp_path)
                        print(f"Đã xóa file tạm: {temp_path}")
                        
                        # Thêm delay giữa các lần xử lý
                        time.sleep(1)
                        
                    else:
                        failed_files.append(file.filename)
                        print(f"File không được hỗ trợ: {file.filename}")
                        continue
                        
                except Exception as e:
                    print(f"Lỗi khi xử lý file {file.filename}: {str(e)}")
                    failed_files.append(file.filename)
                    continue
            
            if not processed_files:
                print("Không có file nào được xử lý thành công")
                return jsonify({'error': 'Không thể xử lý bất kỳ file nào'})
            
            # Lưu tài liệu
            output_filename = f'VBCĐ_{int(time.time())}.docx'
            output_path = os.path.join(temp_dir, output_filename)
            print(f"Đang lưu file Word tại: {output_path}")
            doc.save(output_path)
            
            # Kiểm tra xem file Word có nội dung không
            if os.path.getsize(output_path) == 0:
                print("File Word trống sau khi lưu")
                return jsonify({'error': 'File kết quả trống'})
            
            print(f"File Word đã được lưu thành công, kích thước: {os.path.getsize(output_path)} bytes")
            
            # Tạo file zip chứa kết quả
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'translated_files.zip')
            print(f"Đang tạo file ZIP tại: {zip_path}")
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                zipf.write(output_path, output_filename)
            
            # Read the zip file into memory
            with open(zip_path, 'rb') as f:
                zip_data = f.read()
            
            # Encode the zip data as base64 for JSON transmission
            zip_base64 = base64.b64encode(zip_data).decode('utf-8')
            
            # Xóa thư mục tạm và file zip
            shutil.rmtree(temp_dir)
            os.remove(zip_path)
            print("Đã xóa các file tạm")
            
            return jsonify({
                'success': True,
                'message': 'Dịch thành công!',
                'zip_data': zip_base64,
                'filename': 'translated_files.zip',
                'processed_files': processed_files,
                'failed_files': failed_files
            })
            
        except Exception as e:
            print(f"Lỗi trong quá trình xử lý: {str(e)}")
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            return jsonify({'error': str(e)})
        
    except Exception as e:
        print(f"Lỗi chung: {str(e)}")
        return jsonify({'error': str(e)})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=True)
